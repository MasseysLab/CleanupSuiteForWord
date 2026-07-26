"""Build the focused Word fixture for CleanupSuite structural-safety checks."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Testing Files" / "Regression - Structural Safety.docx"
PRACTICE_OUTPUT = (
    ROOT
    / "Practice - Try CleanupSuite Here"
    / "Testing Files"
    / "Regression - Structural Safety.docx"
)

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF2CC"


def set_font(run, name="Calibri", size=11, bold=False, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def style_table_text(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in paragraph.runs:
                    set_font(run, size=10)


def add_label(doc, title, instruction):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title), bold=True, color="1F4D78")
    p = doc.add_paragraph(instruction)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_font(run, size=10, color="555555")


def add_bookmark(paragraph, bookmark_id, name):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(start)
    paragraph._p.append(end)


def add_empty_field(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " DATE ")
    paragraph._p.append(field)


def add_empty_content_control(paragraph):
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "CleanupSuite protected empty control")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "CleanupSuiteStructuralSafety")
    sdt_pr.extend((alias, tag))
    sdt_content = OxmlElement("w:sdtContent")
    sdt.append(sdt_pr)
    sdt.append(sdt_content)
    paragraph._p.append(sdt)


def build_fixture():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    set_font(title.add_run("CleanupSuite Structural Safety Regression"), size=20, bold=True, color="0B2545")
    subtitle = doc.add_paragraph(
        "Use this fixture to verify that cleanup tools remove only eligible blanks and never silently damage Word table structure or protected content."
    )
    subtitle.paragraph_format.space_after = Pt(10)

    doc.add_heading("1. Blank-paragraph classification", level=1)
    add_label(doc, "Body blank run", "Preview should mark only the second and third blank paragraphs; Apply should leave one blank separator.")
    doc.add_paragraph("BODY BLANK RUN START")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("BODY BLANK RUN END")

    add_label(doc, "NBSP-only paragraph", "The blank-looking paragraph below contains a non-breaking space and must be protected.")
    doc.add_paragraph("NBSP PROTECTION START")
    doc.add_paragraph("\u00a0")
    doc.add_paragraph("NBSP PROTECTION END")

    add_label(doc, "Hidden-text paragraph", "The paragraph below contains hidden text and must be protected.")
    hidden = doc.add_paragraph()
    hidden_run = hidden.add_run("PROTECTED HIDDEN CONTENT")
    set_font(hidden_run)
    hidden_run.font.hidden = True

    add_label(doc, "Bookmark-only paragraph", "The following blank-looking paragraph contains a user bookmark and must be protected.")
    bookmarked = doc.add_paragraph()
    add_bookmark(bookmarked, 41, "CleanupSuiteProtectedBlank")

    add_label(doc, "Field-only paragraph", "The following blank-looking paragraph contains a DATE field and must be protected.")
    field_paragraph = doc.add_paragraph()
    add_empty_field(field_paragraph)

    add_label(doc, "Content-control paragraph", "The following blank-looking paragraph contains an empty content control and must be protected.")
    control_paragraph = doc.add_paragraph()
    add_empty_content_control(control_paragraph)

    doc.add_heading("2. Supported table candidates", level=1)
    add_label(doc, "Empty row candidate", "Only the middle row is eligible. The header and populated row must remain.")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Protected header A"
    table.cell(0, 1).text = "Protected header B"
    mark_header_row(table.rows[0])
    shade_cell(table.cell(0, 0), PALE_BLUE)
    shade_cell(table.cell(0, 1), PALE_BLUE)
    table.cell(2, 0).text = "Keep this populated row"
    table.cell(2, 1).text = "Content"
    set_table_geometry(table, [4680, 4680])
    style_table_text(table)

    add_label(doc, "Empty column candidate", "Only the center column is eligible; the two populated columns must remain.")
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Keep left"
    table.cell(0, 2).text = "Keep right"
    table.cell(1, 0).text = "Left content"
    table.cell(1, 2).text = "Right content"
    set_table_geometry(table, [3120, 3120, 3120])
    style_table_text(table)

    add_label(doc, "Extra blank paragraphs inside a cell", "Paragraph Fixer may collapse extras, but the required cell and row markers must remain.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).add_paragraph("")
    table.cell(0, 0).add_paragraph("")
    table.cell(0, 1).text = "Neighbor cell must remain"
    set_table_geometry(table, [4680, 4680])
    style_table_text(table)

    doc.add_heading("3. Protected or unsupported table structures", level=1)
    add_label(doc, "Blank repeating header", "The blank first row is a header row and must never be deleted.")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    shade_cell(table.cell(0, 0), PALE_GOLD)
    shade_cell(table.cell(0, 1), PALE_GOLD)
    table.cell(1, 0).text = "Content row"
    table.cell(1, 1).text = "Must remain"
    set_table_geometry(table, [4680, 4680])
    style_table_text(table)

    add_label(doc, "Merged table", "This nonuniform table must be reported as protected/unsupported and left unchanged.")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged heading"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = "Keep"
    set_table_geometry(table, [4680, 4680])
    style_table_text(table)

    add_label(doc, "Nested table", "The outer table contains a nested table and must be reported as protected/unsupported.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Outer cell"
    nested = table.cell(0, 1).add_table(rows=1, cols=1)
    nested.style = "Table Grid"
    nested.cell(0, 0).text = "Nested content"
    set_table_geometry(nested, [4200])
    set_table_geometry(table, [4680, 4680])
    style_table_text(table)

    doc.add_heading("4. Table-to-text parity", level=1)
    add_label(doc, "Eligible single-column table", "Preview and Apply should both count this table for conversion.")
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "First line"
    table.cell(1, 0).text = "Second line"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    style_table_text(table)

    add_label(doc, "Protected multi-column conversion", "Preview and Apply must both skip this table when Convert to text is selected.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Left"
    table.cell(0, 1).text = "Right"
    set_table_geometry(table, [4680, 4680])
    style_table_text(table)

    doc.add_heading("5. Final-document protection", level=1)
    doc.add_paragraph("The document ends with two removable body blanks followed by Word's required final paragraph marker.")
    doc.add_paragraph("FINAL CONTENT SENTINEL")
    doc.add_paragraph("")
    doc.add_paragraph("")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    PRACTICE_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    copy2(OUTPUT, PRACTICE_OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_fixture())
