from dataclasses import dataclass
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import shutil
import subprocess
import tempfile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ROOT_TESTS = ROOT / "Testing Files"
PRACTICE_TESTS = ROOT / "Practice - Try CleanupSuite Here" / "Testing Files"
TAN = "FFF2CC"
GREEN = "E2F0D9"


@dataclass(frozen=True)
class Case:
    label: str
    before: str
    after: str
    kind: str = "text"


@dataclass(frozen=True)
class Tool:
    filename: str
    name: str
    cases: tuple[Case, ...]


def c(label, before, after, kind="text"):
    return Case(label, f"{label} | {before}", f"{label} | {after}", kind)


TOOLS = (
    Tool("TTF-01_InvisibleUnicodeCleaner.docx", "Invisible Unicode Cleaner", (
        c("Non-breaking spaces", "Alpha\u00a0Bravo\u00a0Charlie", "Alpha Bravo Charlie"),
        c("Zero-width spaces", "zero\u200bwidth\u200bcharacters", "zerowidthcharacters"),
        c("Zero-width non-joiners", "non\u200cjoiner", "nonjoiner"),
        c("Zero-width joiners", "joined\u200dword", "joinedword"),
        c("Byte order marks", "hidden\ufeffmark", "hiddenmark"),
        c("Soft hyphens", "soft\u00adhyphen", "soft-hyphen"),
        c("Non-breaking hyphens", "non\u2011breaking\u2011hyphen", "non-breaking-hyphen"),
    )),
    Tool("TTF-02_PunctuationNormalizer.docx", "Punctuation Normalizer", (
        c("Curly double quotes", "\u201cQuoted words remain readable.\u201d", '"Quoted words remain readable."'),
        c("Curly single quotes", "It\u2019s copied from Word\u2019s editor.", "It's copied from Word's editor."),
        c("Em dashes", "A pause\u2014then the sentence continues.", "A pause-then the sentence continues."),
        c("En dashes", "Pages 10\u201315 contain the example.", "Pages 10-15 contain the example."),
        c("Ellipses", "Wait\u2026the sentence continues.", "Wait...the sentence continues."),
    )),
    Tool("TTF-03_SpacingFixer.docx", "Spacing Fixer", (
        c("Multiple spaces", "This  sentence   contains    repeated spaces.", "This sentence contains repeated spaces."),
        c("Leading and trailing spaces", "   The sentence is padded.   ", "The sentence is padded."),
        c("Spaces before punctuation", "Hello , world . Why ? Yes !", "Hello, world. Why? Yes!"),
        c("Spacing after punctuation", "First.  Second,  third:  fourth.", "First. Second, third: fourth."),
        c("Extra blank lines", "Two content lines use paragraph spacing.", "Two content lines use paragraph spacing.", "blank_lines"),
    )),
    Tool("TTF-04_CapitalizationFixer.docx", "Capitalization Fixer", (
        c("Sentence boundaries", "the first sentence starts low. the second sentence starts low.", "The first sentence starts low. The second sentence starts low."),
        c("Abbreviations", "dr. smith met mr. jones before noon.", "Dr. smith met Mr. jones before noon."),
        c("Acronyms", "the nasa api exports pdf reports for vba users.", "The NASA API exports PDF reports for VBA users."),
        c("Names and brands", "cleanupsuite works with github and chatgpt.", "CleanupSuite works with GitHub and ChatGPT."),
        c("Heading context", "a pasted heading for the document", "A Pasted Heading for the Document", "capitalization_heading"),
    )),
    Tool("TTF-05_ListNormalizer.docx", "List Normalizer", (
        c("Manual hyphen bullets", "First item in the list", "First item in the list", "list_hyphen"),
        c("Manual asterisk bullets", "Second item in the list", "Second item in the list", "list_asterisk"),
        c("Manual numbered items", "First numbered instruction", "First numbered instruction", "list_number"),
        c("List indentation", "A real list item", "A real list item", "list_indent"),
    )),
    Tool("TTF-06_ParagraphStructureFixer.docx", "Paragraph Structure Fixer", (
        c("Consecutive empty paragraphs", "Two content lines use paragraph spacing.", "Two content lines use paragraph spacing.", "blank_lines"),
        c("Soft returns to paragraph marks", "First line / Second line / Third line", "First line / Second line / Third line", "soft_to_para"),
        c("Paragraph spacing", "This paragraph uses extra Space Before and Space After.", "This paragraph uses extra Space Before and Space After.", "paragraph_spacing"),
        c("Manual tab indentation", "This line starts with a manual tab.", "This line starts with a manual tab.", "manual_tab_indent"),
        c("Manual space indentation", "This line starts with manual spaces.", "This line starts with manual spaces.", "manual_space_indent"),
    )),
    Tool("TTF-07_SoftReturnConverter.docx", "Soft Return Converter", (
        c("Soft returns to paragraph marks", "First line / Second line / Third line", "First line / Second line / Third line", "soft_to_para"),
        c("Paragraph marks to soft returns", "First line / Second line / Third line", "First line / Second line / Third line", "para_to_soft"),
    )),
    Tool("TTF-08_TrimEndOfDoc.docx", "Trim End of Doc", (
        c("Trailing empty paragraphs", "The document ends after this sentence.", "The document ends after this sentence.", "trailing_blanks"),
    )),
    Tool("TTF-09_TableCleaner.docx", "Table Cleaner", tuple(
        c(label, "Shared table sample.", "Shared table sample.", "table")
        for label in ("Empty table rows", "Empty table columns", "Cell padding", "Direct cell formatting", "Border normalization", "Border removal", "Single-column table to text")
    )),
    Tool("TTF-10_HeaderFooterStandardizer.docx", "Header / Footer Standardizer", tuple(
        c(label, "Shared header and footer sample.", "Shared header and footer sample.")
        for label in ("Header formatting", "Footer formatting", "Header and footer font", "Header and footer spacing", "Header and footer alignment", "Unlink sections", "Clear header and footer content")
    )),
    Tool("TTF-11_FootnoteEndnoteRemover.docx", "Footnote / Endnote Remover", (
        c("Remove footnotes", "The sentence has a footnote reference.", "The sentence has a footnote reference.", "footnote"),
        c("Remove endnotes", "The sentence has an endnote reference.", "The sentence has an endnote reference.", "endnote"),
        c("Keep note text inline", "The sentence keeps its note.", "The sentence keeps its note [kept note text].", "keep_note"),
    )),
    Tool("TTF-12_ObjectRemover.docx", "Object Remover", (
        c("Pictures", "The sample contains a removable picture.", "The sample contains a removable picture.", "picture"),
        c("Text boxes", "The sample contains a removable text box.", "The sample contains a removable text box."),
        c("Frames", "The sample contains a removable frame around text.", "The sample contains a removable frame around text.", "frame"),
        c("Horizontal lines", "The sample contains a removable horizontal line.", "The sample contains a removable horizontal line."),
        c("Form controls", "The sample contains a removable form control.", "The sample contains a removable form control."),
        c("Hidden text", "Visible text surrounds removable hidden text.", "Visible text surrounds removable hidden text.", "hidden"),
        c("Tables", "The sample contains a removable table.", "The sample contains a removable table.", "object_table"),
    )),
    Tool("TTF-13_FontNormalizer.docx", "Font Normalizer", (
        c("Font face", "Shared font sample.", "Shared font sample.", "font_face"),
        c("Font size", "Shared font sample.", "Shared font sample.", "font_size"),
        c("Direct bold", "Shared font sample.", "Shared font sample.", "bold"),
        c("Direct italic", "Shared font sample.", "Shared font sample.", "italic"),
        c("Font color", "Shared font sample.", "Shared font sample.", "font_color"),
    )),
    Tool("TTF-14_FormattingCleaner.docx", "Formatting Cleaner", (
        c("Direct character formatting", "Shared formatting sample.", "Shared formatting sample.", "bold"),
        c("Direct paragraph formatting", "Shared formatting sample.", "Shared formatting sample.", "paragraph_spacing"),
        c("Preserve highlighting", "The emphasized words remain highlighted.", "The emphasized words remain highlighted.", "highlight"),
        c("Quick emphasis", "A whole paragraph keeps its emphasis.", "A whole paragraph keeps its emphasis.", "bold_both"),
        c("Thorough emphasis", "One important word keeps emphasis.", "One important word keeps emphasis.", "italic_both"),
        c("Strip emphasis", "Direct emphasis is removed from this sentence.", "Direct emphasis is removed from this sentence.", "bold"),
    )),
    Tool("TTF-15_DuplicateParagraphDetector.docx", "Duplicate Paragraph Detector", (
        c("Exact duplicate removal", "Shared duplicate paragraph.", "Shared duplicate paragraph.", "duplicate"),
        c("Normalized duplicate removal", "Shared normalized duplicate paragraph.", "Shared normalized duplicate paragraph.", "duplicate"),
        c("Fuzzy duplicate removal", "Shared fuzzy duplicate paragraph.", "Shared fuzzy duplicate paragraph.", "duplicate"),
    )),
    Tool("TTF-16_MetaDataSuite.docx", "MetaDataSuite", tuple(
        c(label, "Shared metadata sample.", "Shared metadata sample.")
        for label in ("Document properties", "Personal information", "Sharing properties", "Author identity")
    )),
    Tool("TTF-17_FinalReview.docx", "Final Review", tuple(
        c(label, "Shared review sentence.", "Shared review sentence.")
        for label in ("Remove comments", "Accept tracked changes", "Minimal author cleanup", "Normal author cleanup", "Broad author cleanup")
    )),
    Tool("TTF-18_HyperlinkCleaner.docx", "Hyperlink Cleaner", (
        c("Hide hyperlink styling", "Visit CleanupSuite documentation.", "Visit CleanupSuite documentation.", "hyperlink"),
        c("Remove hyperlinks", "Open the CleanupSuite reference page.", "Open the CleanupSuite reference page.", "hyperlink"),
    )),
    Tool("TTF-19_StyleCleanup.docx", "Style Cleanup", (
        c("Unused custom styles", "Shared style sample.", "Shared style sample."),
        c("Style variant remapping", "Shared style sample.", "Shared style sample.", "variant_style"),
    )),
    Tool("TTF-20_BreakNormalizer.docx", "Break Normalizer", (
        c("Consecutive section breaks", "Shared section boundary sample.", "Shared section boundary sample.", "section_break"),
        c("Next Page conversion", "Shared section conversion sample.", "Shared section conversion sample."),
        c("Continuous conversion", "Shared section conversion sample.", "Shared section conversion sample."),
        c("Even Page conversion", "Shared section conversion sample.", "Shared section conversion sample."),
        c("Odd Page conversion", "Shared section conversion sample.", "Shared section conversion sample."),
        c("Consecutive page breaks", "Shared manual page break sample.", "Shared manual page break sample.", "page_break"),
    )),
)


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_sample(paragraph, fill, bottom=1):
    shade(paragraph, fill)
    paragraph.paragraph_format.keep_together = True


def add_sample_paragraph(doc, text, fill):
    paragraph = doc.add_paragraph(text)
    format_sample(paragraph, fill)
    return paragraph


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    r_pr.append(style)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_table_borders(table, enabled):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single" if enabled else "nil")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), "808080")


def set_table_padding(table, horizontal_twips, vertical_twips):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for edge_name, width in (("top", vertical_twips), ("left", horizontal_twips), ("bottom", vertical_twips), ("right", horizontal_twips)):
        edge = margins.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            margins.append(edge)
        edge.set(qn("w:w"), str(width))
        edge.set(qn("w:type"), "dxa")


def add_table_sample(doc, case, fill, dirty):
    add_sample_paragraph(doc, case.before if dirty else case.after, fill)
    rows, cols = (3, 2)
    if case.label == "Empty table columns":
        cols = 3 if dirty else 2
    if case.label == "Empty table rows":
        rows = 3 if dirty else 2
    if case.label == "Single-column table to text" and not dirty:
        add_sample_paragraph(doc, "Alpha", fill)
        add_sample_paragraph(doc, "Beta", fill)
        return
    if case.label == "Single-column table to text":
        rows = 2
        cols = 1
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            empty_row_target = dirty and case.label == "Empty table rows" and row_index == rows - 1
            empty_col_target = dirty and case.label == "Empty table columns" and col_index == cols - 1
            if case.label == "Single-column table to text":
                cell.text = ("Alpha", "Beta")[row_index]
            else:
                cell.text = "" if empty_row_target or empty_col_target else f"Cell {row_index + 1}.{col_index + 1}"
            if case.label == "Direct cell formatting" and dirty:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 0, 0)
    if case.label == "Cell padding":
        set_table_padding(table, 288 if dirty else 115, 288 if dirty else 58)
    if case.label == "Border normalization":
        set_table_borders(table, not dirty)
    if case.label == "Border removal":
        set_table_borders(table, dirty)


def add_case(doc, case, target_image):
    kind = case.kind
    if kind == "blank_lines":
        add_sample_paragraph(doc, case.before, TAN)
        for _ in range(3):
            add_sample_paragraph(doc, "", TAN)
        add_sample_paragraph(doc, "Second content line.", TAN)
        add_sample_paragraph(doc, case.after, GREEN)
        add_sample_paragraph(doc, "", GREEN)
        add_sample_paragraph(doc, "Second content line.", GREEN)
    elif kind in {"soft_to_para", "para_to_soft"}:
        if kind == "soft_to_para":
            tan = doc.add_paragraph()
            tan.add_run(case.label + " | First line")
            tan.add_run().add_break(WD_BREAK.LINE)
            tan.add_run("Second line")
            tan.add_run().add_break(WD_BREAK.LINE)
            tan.add_run("Third line")
            format_sample(tan, TAN)
            add_sample_paragraph(doc, case.label + " | First line", GREEN)
            add_sample_paragraph(doc, "Second line", GREEN)
            add_sample_paragraph(doc, "Third line", GREEN)
        else:
            add_sample_paragraph(doc, case.label + " | First line", TAN)
            add_sample_paragraph(doc, "Second line", TAN)
            add_sample_paragraph(doc, "Third line", TAN)
            green = doc.add_paragraph()
            green.add_run(case.label + " | First line")
            green.add_run().add_break(WD_BREAK.LINE)
            green.add_run("Second line")
            green.add_run().add_break(WD_BREAK.LINE)
            green.add_run("Third line")
            format_sample(green, GREEN)
    elif kind.startswith("list_"):
        if kind == "list_hyphen":
            add_sample_paragraph(doc, "- " + case.before, TAN)
        elif kind == "list_asterisk":
            add_sample_paragraph(doc, "* " + case.before, TAN)
        else:
            tan = doc.add_paragraph(case.before if kind == "list_indent" else "1. " + case.before)
            tan.style = "List Bullet" if kind == "list_indent" else doc.styles["Normal"]
            if kind == "list_indent":
                tan.paragraph_format.left_indent = Inches(1.2)
            format_sample(tan, TAN)
        green = doc.add_paragraph(case.after)
        green.style = "List Number" if kind == "list_number" else "List Bullet"
        format_sample(green, GREEN)
    elif kind == "paragraph_spacing":
        tan = add_sample_paragraph(doc, case.before, TAN)
        tan.paragraph_format.space_before = Pt(18)
        tan.paragraph_format.space_after = Pt(18)
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind in {"manual_tab_indent", "manual_space_indent"}:
        prefix = "\t" if kind == "manual_tab_indent" else "    "
        add_sample_paragraph(doc, prefix + case.before, TAN)
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "table":
        add_table_sample(doc, case, TAN, True)
        add_table_sample(doc, case, GREEN, False)
    elif kind in {"font_face", "font_size", "bold", "italic", "font_color", "highlight", "bold_both", "italic_both"}:
        tan = add_sample_paragraph(doc, case.before, TAN)
        green = add_sample_paragraph(doc, case.after, GREEN)
        if kind == "font_face":
            tan.runs[0].font.name = "Courier New"
        elif kind == "font_size":
            tan.runs[0].font.size = Pt(18)
        elif kind in {"bold", "bold_both"}:
            tan.runs[0].bold = True
            if kind == "bold_both": green.runs[0].bold = True
        elif kind in {"italic", "italic_both"}:
            tan.runs[0].italic = True
            if kind == "italic_both": green.runs[0].italic = True
        elif kind == "font_color":
            tan.runs[0].font.color.rgb = RGBColor(192, 0, 0)
        elif kind == "highlight":
            tan.runs[0].font.highlight_color = 7
            green.runs[0].font.highlight_color = 7
    elif kind == "duplicate":
        add_sample_paragraph(doc, case.before, TAN)
        add_sample_paragraph(doc, case.before, TAN)
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "hyperlink":
        tan = doc.add_paragraph(case.label + " | ")
        add_hyperlink(tan, case.before.split(" | ", 1)[1], "https://example.com/cleanupsuite")
        format_sample(tan, TAN)
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "picture":
        tan = add_sample_paragraph(doc, case.before + " ", TAN)
        tan.add_run().add_picture(str(target_image), width=Inches(0.25))
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "hidden":
        tan = add_sample_paragraph(doc, case.before + " ", TAN)
        hidden = tan.add_run("REMOVE THIS HIDDEN TEXT")
        hidden.font.hidden = True
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "capitalization_heading":
        tan = add_sample_paragraph(doc, case.before, TAN)
        tan.style = "Heading 1"
        green = add_sample_paragraph(doc, case.after, GREEN)
        green.style = "Heading 1"
    elif kind == "frame":
        tan = add_sample_paragraph(doc, case.before, TAN)
        frame_pr = OxmlElement("w:framePr")
        frame_pr.set(qn("w:w"), "5000")
        tan._p.get_or_add_pPr().append(frame_pr)
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "object_table":
        add_sample_paragraph(doc, case.before, TAN)
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "REMOVABLE TABLE TARGET"
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "variant_style":
        tan = add_sample_paragraph(doc, case.before, TAN)
        tan.style = "Body Text 2"
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "page_break":
        tan = add_sample_paragraph(doc, case.before, TAN)
        tan.add_run().add_break(WD_BREAK.PAGE)
        tan.add_run().add_break(WD_BREAK.PAGE)
        add_sample_paragraph(doc, case.after, GREEN)
    elif kind == "section_break":
        add_sample_paragraph(doc, case.before, TAN)
        doc.add_section(WD_SECTION.CONTINUOUS)
        doc.add_section(WD_SECTION.CONTINUOUS)
        add_sample_paragraph(doc, case.after, GREEN)
        doc.add_section(WD_SECTION.CONTINUOUS)
    else:
        add_sample_paragraph(doc, case.before, TAN)
        add_sample_paragraph(doc, case.after, GREEN)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def create_legend(path):
    image = Image.new("RGB", (1200, 170), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 30)
        bold = ImageFont.truetype("arialbd.ttf", 31)
    except OSError:
        font = ImageFont.load_default()
        bold = font
    rows = [
        ((255, 242, 204), "Tan:", "This is how the test looks before the selected action is applied."),
        ((226, 240, 217), "Green:", "The tan sample should look like this after the selected action is applied."),
    ]
    for index, (color, label, text) in enumerate(rows):
        top = 8 + index * 80
        draw.rounded_rectangle((8, top, 1192, top + 68), radius=8, fill=color, outline=(170, 170, 160), width=2)
        draw.text((28, top + 16), label, fill=(35, 35, 35), font=bold)
        label_width = draw.textlength(label, font=bold)
        draw.text((45 + label_width, top + 16), text, fill=(35, 35, 35), font=font)
    image.save(path)


def create_title_image(path, title_text):
    image = Image.new("RGB", (1200, 72), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arialbd.ttf", 34)
    except OSError:
        font = ImageFont.load_default()
    draw.text((8, 14), title_text, fill=(20, 32, 45), font=font)
    image.save(path)


def add_title_and_legend(doc, index, tool, legend_path, title_image_path, use_header=False):
    container = doc.sections[0].header if use_header else doc
    title = container.paragraphs[0] if use_header else container.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title_text = f"TTF-{index:02d} | {tool.name}"
    if use_header:
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(16)
    else:
        inline = title.add_run().add_picture(str(title_image_path), width=Inches(6.65))
        inline._inline.docPr.set("descr", title_text)
    picture = container.add_paragraph()
    picture.paragraph_format.space_after = Pt(10)
    picture.add_run().add_picture(str(legend_path), width=Inches(6.65))


def replace_zip_part(path, updates, additions=None):
    additions = additions or {}
    temp_path = path.with_suffix(path.suffix + ".tmp")
    with ZipFile(path, "r") as src, ZipFile(temp_path, "w", ZIP_DEFLATED) as dst:
        written = set()
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename in updates:
                data = updates[item.filename](data)
            dst.writestr(item, data)
            written.add(item.filename)
        for name, data in additions.items():
            if name not in written:
                dst.writestr(name, data)
    temp_path.replace(path)


def ensure_content_type(xml, part_name, content_type):
    text = xml.decode("utf-8")
    if part_name not in text:
        text = text.replace("</Types>", f'<Override PartName="/{part_name}" ContentType="{content_type}"/></Types>')
    return text.encode("utf-8")


def ensure_relationship(xml, rel_id, rel_type, target):
    text = xml.decode("utf-8")
    if rel_id not in text:
        text = text.replace("</Relationships>", f'<Relationship Id="{rel_id}" Type="{rel_type}" Target="{target}"/></Relationships>')
    return text.encode("utf-8")


def insert_reference_after_marker(text, marker, reference_xml):
    marker_index = text.find(marker)
    if marker_index < 0:
        raise RuntimeError(f"Fixture marker not found: {marker}")
    paragraph_end = text.find("</w:p>", marker_index)
    return text[:paragraph_end] + reference_xml + text[paragraph_end:]


def patch_notes_fixture(path):
    def update_document(xml):
        text = xml.decode("utf-8")
        text = insert_reference_after_marker(text, "Remove footnotes", '<w:r><w:footnoteReference w:id="2"/></w:r>')
        text = insert_reference_after_marker(text, "Remove endnotes", '<w:r><w:endnoteReference w:id="2"/></w:r>')
        text = insert_reference_after_marker(text, "Keep note text inline", '<w:r><w:footnoteReference w:id="3"/></w:r>')
        return text.encode("utf-8")

    footnotes = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
        '<w:footnote w:id="2"><w:p><w:r><w:footnoteRef/></w:r><w:r><w:t>removable footnote text</w:t></w:r></w:p></w:footnote>'
        '<w:footnote w:id="3"><w:p><w:r><w:footnoteRef/></w:r><w:r><w:t>kept note text</w:t></w:r></w:p></w:footnote></w:footnotes>').encode("utf-8")
    endnotes = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:endnote>'
        '<w:endnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>'
        '<w:endnote w:id="2"><w:p><w:r><w:endnoteRef/></w:r><w:r><w:t>removable endnote text</w:t></w:r></w:p></w:endnote></w:endnotes>').encode("utf-8")
    replace_zip_part(path, {
        "word/document.xml": update_document,
        "word/_rels/document.xml.rels": lambda xml: ensure_relationship(ensure_relationship(xml, "rIdFixtureFootnotes", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes", "footnotes.xml"), "rIdFixtureEndnotes", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes", "endnotes.xml"),
        "[Content_Types].xml": lambda xml: ensure_content_type(ensure_content_type(xml, "word/footnotes.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"), "word/endnotes.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"),
    }, {"word/footnotes.xml": footnotes, "word/endnotes.xml": endnotes})


def patch_review_fixture(path):
    def update_document(xml):
        text = xml.decode("utf-8")
        marker = "Remove comments"
        marker_index = text.find(marker)
        run_start = text.rfind("<w:r", 0, marker_index)
        run_end = text.find("</w:r>", marker_index) + len("</w:r>")
        text = text[:run_start] + '<w:commentRangeStart w:id="0"/>' + text[run_start:run_end] + '<w:commentRangeEnd w:id="0"/><w:r><w:commentReference w:id="0"/></w:r>' + text[run_end:]
        marker = "Accept tracked changes"
        marker_index = text.find(marker)
        run_start = text.rfind("<w:r", 0, marker_index)
        run_end = text.find("</w:r>", marker_index) + len("</w:r>")
        text = text[:run_start] + '<w:ins w:id="1" w:author="Fixture Author" w:date="2026-07-10T00:00:00Z">' + text[run_start:run_end] + '</w:ins>' + text[run_end:]
        return text.encode("utf-8")

    comments = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:comment w:id="0" w:author="Fixture Author" w:date="2026-07-10T00:00:00Z"><w:p><w:r><w:t>removable fixture comment</w:t></w:r></w:p></w:comment>'
        '</w:comments>').encode("utf-8")
    replace_zip_part(path, {
        "word/document.xml": update_document,
        "word/_rels/document.xml.rels": lambda xml: ensure_relationship(xml, "rIdFixtureComments", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments", "comments.xml"),
        "[Content_Types].xml": lambda xml: ensure_content_type(xml, "word/comments.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"),
    }, {"word/comments.xml": comments})


def configure_document(doc, tool):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    if "Body Text 2" not in [style.name for style in styles]:
        styles.add_style("Body Text 2", WD_STYLE_TYPE.PARAGRAPH)
    if tool.name == "MetaDataSuite":
        props = doc.core_properties
        props.title = "Fixture Document Title"
        props.subject = "Fixture Subject"
        props.author = "Fixture Author"
        props.keywords = "fixture, cleanup, metadata"
        props.comments = "Fixture sharing metadata"
    if tool.name == "Style Cleanup":
        styles.add_style("Unused Fixture Style", WD_STYLE_TYPE.PARAGRAPH)


def build_fixture(index, tool, legend_path, target_image, title_image_path):
    doc = Document()
    configure_document(doc, tool)
    add_title_and_legend(doc, index, tool, legend_path, title_image_path, use_header=(tool.name == "Object Remover"))
    for case in tool.cases:
        add_case(doc, case, target_image)
    if tool.name == "Trim End of Doc":
        for _ in range(7):
            doc.add_paragraph("")
    if tool.name == "Header / Footer Standardizer":
        header = doc.sections[0].header.paragraphs[0]
        header.text = "Header formatting | Shared header sample"
        header.runs[0].font.name = "Courier New"
        header.paragraph_format.space_after = Pt(18)
        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = "Footer formatting | Shared footer sample"
        footer.runs[0].font.name = "Times New Roman"
        footer.paragraph_format.space_before = Pt(18)
        doc.add_section(WD_SECTION.CONTINUOUS)
    out = ROOT_TESTS / tool.filename
    doc.save(out)
    if tool.name == "Footnote / Endnote Remover":
        patch_notes_fixture(out)
    if tool.name == "Final Review":
        patch_review_fixture(out)
    if tool.name == "Object Remover":
        subprocess.run([
            "powershell",
            "-ExecutionPolicy", "Bypass",
            "-File", str(ROOT / "scripts" / "enrich_object_fixture.ps1"),
            "-DocumentPath", str(out),
        ], check=True)
    shutil.copyfile(out, PRACTICE_TESTS / tool.filename)


def purge_old_fixtures():
    root_resolved = ROOT.resolve()
    for path in ROOT.rglob("TTF-*.docx"):
        resolved = path.resolve()
        if root_resolved not in resolved.parents:
            raise RuntimeError(f"Refusing to delete fixture outside workspace: {resolved}")
        path.unlink()


def main():
    purge_old_fixtures()
    ROOT_TESTS.mkdir(parents=True, exist_ok=True)
    PRACTICE_TESTS.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        legend_path = temp_dir / "fixture-legend.png"
        target_image = temp_dir / "removable-picture.png"
        create_legend(legend_path)
        Image.new("RGB", (120, 80), (60, 130, 210)).save(target_image)
        for index, tool in enumerate(TOOLS, start=1):
            title_image_path = temp_dir / f"fixture-title-{index:02d}.png"
            create_title_image(title_image_path, f"TTF-{index:02d} | {tool.name}")
            build_fixture(index, tool, legend_path, target_image, title_image_path)


if __name__ == "__main__":
    main()
