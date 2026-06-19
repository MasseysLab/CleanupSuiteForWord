from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)


def add_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        else:
            paragraph.add_run(part)


def convert_markdown(markdown_path: Path, docx_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_defaults(doc)

    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        text = " ".join(line.strip() for line in paragraph_lines).strip()
        if text:
            p = doc.add_paragraph()
            add_runs(p, text)
        paragraph_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        for line in code_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(9)
        code_lines = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        if line.strip().startswith("```"):
            flush_paragraph()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_paragraph()
            continue

        if line.startswith("# "):
            flush_paragraph()
            doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:].strip(), level=2)
            continue

        if line.startswith("> "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.style = "Intense Quote" if "Intense Quote" in doc.styles else "Quote"
            add_runs(p, line[2:].strip())
            continue

        bullet_match = re.match(r"^[-*] (.+)$", line)
        if bullet_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, bullet_match.group(1).strip())
            continue

        number_match = re.match(r"^\d+\. (.+)$", line)
        if number_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_runs(p, number_match.group(1).strip())
            continue

        if re.match(r"^---+$", line.strip()):
            flush_paragraph()
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    flush_code()

    doc.core_properties.title = markdown_path.stem.replace("_", " ")
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("usage: export_markdown_docs_to_docx.py <input.md> <output.docx>")
        return 1
    src = (ROOT / argv[1]).resolve() if not Path(argv[1]).is_absolute() else Path(argv[1])
    dst = (ROOT / argv[2]).resolve() if not Path(argv[2]).is_absolute() else Path(argv[2])
    convert_markdown(src, dst)
    print(f"Wrote {dst}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
